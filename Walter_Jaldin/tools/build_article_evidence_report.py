from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACK = ROOT / "evidencias" / "publicable"
OUT_DIR = ROOT / "08_articulo" / "entregables"
OUT_DOCX = OUT_DIR / "Articulo_Evidencias_FutbolLibre_DoradoBet.docx"
OUT_MD = OUT_DIR / "Articulo_Evidencias_FutbolLibre_DoradoBet.md"


def read_csv(name: str) -> list[dict[str, str]]:
    with (PACK / name).open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def paragraph(doc: Document, text: str, *, bold_label: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        rest = p.add_run(text[len(bold_label):])
        rest.font.name = "Calibri"
        rest.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level < 3 else 120)
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    return p


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(9)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(8.5)
    if widths:
        for row in tbl.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return tbl


def add_figure(doc: Document, image_path: Path, caption: str):
    if not image_path.exists():
        paragraph(doc, f"No se encontro la captura esperada: {image_path}")
        return
    doc.add_picture(str(image_path), width=Inches(6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(9)


scenario_rows = read_csv("scenario_summary.csv")
clone_rows = read_csv("doradobet_clone_scan.csv")

scenario_lookup = {row["scenario"]: row for row in scenario_rows}

proofs = [
    {
        "id": "P1",
        "title": "Futbol Libre carga publicidad ofuscada desde terceros",
        "finding": "La portada de Futbol Libre carga acscdn.com, adexchangerapid.com y usrpubtrk.com. Tambien se conservo copia local de los scripts aclib.js y suv5.js.",
        "evidence": "evidencias/aclib.js; evidencias/suv5.js; evidencias/publicable/futbol_home_desktop.redacted.json",
        "risk": "Malvertising, rastreo y redirecciones variables por campana.",
    },
    {
        "id": "P2",
        "title": "Click en el reproductor intenta abrir ventana nueva",
        "finding": "En escritorio, el escenario futbol_espn_desktop_click registro 1 popup desde el iframe del reproductor. No hubo descarga.",
        "evidence": "evidencias/extended/futbol_espn_desktop_click/probe-result.json; captura after.png",
        "risk": "Pop-under y landing externa potencialmente cambiante.",
    },
    {
        "id": "P3",
        "title": "El stream usa iframe externo y hosts rotativos",
        "finding": "El canal ESPN se embebe desde latamvidz1.com y contacta hosts envivoslatam.org distintos entre corridas, ademas de infraestructura P2P/telemetria.",
        "evidence": "evidencias/latamvidz1_espn.html; dominios ng0pr.envivoslatam.org y smjt9q.envivoslatam.org en JSON redactados.",
        "risk": "Dependencia de terceros y comportamiento cambiante fuera del dominio principal.",
    },
    {
        "id": "P4",
        "title": "DoradoBet movil expone una cadena adult/affiliate",
        "finding": "En mobile/PWA se observaron solicitudes hacia bit.ly, go.xlviiirdr.com y chaturbate.com con track=doradobet.com.",
        "evidence": "evidencias/publicable/interesting_requests.csv; doradobet_home_mobile.redacted.json; doradobet_pwa_mobile.redacted.json",
        "risk": "Redireccion/publicidad adulta o afiliada asociada al trafico movil, aunque sin descarga observada.",
    },
    {
        "id": "P5",
        "title": "DoradoBet tiene rastreo intensivo",
        "finding": "La carga desktop contacto 62 dominios; mobile 86; registro 69; PWA 68. Incluye FingerprintJS, Amplitude, TikTok, Clarity, Criteo, Adform, Xandr/AppNexus, MGID y Sportradar.",
        "evidence": "evidencias/publicable/scenario_summary.csv; domains_by_scenario.csv",
        "risk": "Perfilado, retargeting y huella persistente del visitante.",
    },
    {
        "id": "P6",
        "title": "Fingerprinting confirmado por codigo",
        "finding": "El HTML de DoradoBet carga fpjscdn.net, guarda visitor_fp/visitor_fp5/fp_visitor_id y envia datos a static.virtualsoft.tech/setvid.",
        "evidence": "evidencias/doradobet/home.html; doradobet_home_desktop.redacted.json",
        "risk": "Identificacion persistente incluso sin iniciar sesion.",
    },
    {
        "id": "P7",
        "title": "Registro pide datos personales sensibles",
        "finding": "La pantalla de registro solicita nombre, apellido, documento, numero de identificacion, telefono, ciudad/provincia, email y contrasena.",
        "evidence": "evidencias/extended/doradobet_register_desktop/after.png",
        "risk": "Riesgo KYC/identidad, especialmente si el usuario cae en un clon.",
    },
    {
        "id": "P8",
        "title": "La PWA carga infraestructura de engagement/push",
        "finding": "La ruta /landing/app-pwa contacto Optimove y Kumulos, incluidos endpoints de worker/config/eventos. No se observo prompt de notificaciones.",
        "evidence": "evidencias/publicable/doradobet_pwa_mobile.redacted.json; interesting_requests.csv",
        "risk": "Seguimiento de instalacion/eventos y posibilidad de engagement push segun contexto.",
    },
    {
        "id": "P9",
        "title": "Dominios parecidos amplian el riesgo de suplantacion",
        "finding": "El escaneo encontro dominios similares con Cloudflare, WordPress, LiteSpeed, titulos promocionales y un caso con X-Powered-By: PHP/5.6.40.",
        "evidence": "evidencias/publicable/doradobet_clone_scan.csv",
        "risk": "Confusion de marca, phishing, captura de datos o afiliacion agresiva.",
    },
    {
        "id": "P10",
        "title": "No hubo descarga automatica en los escenarios medidos",
        "finding": "Los siete escenarios registraron downloads=0. Esto limita la conclusion: no se probo malware directo, pero si redirecciones, popups y rastreo.",
        "evidence": "evidencias/publicable/scenario_summary.csv",
        "risk": "La ausencia de descarga en una muestra no descarta campanas futuras o segmentadas.",
    },
]


scenario_table_rows = [
    [
        row["scenario"],
        row["domains_count"],
        row["popups"],
        row["downloads"],
        row["local_storage_keys"],
    ]
    for row in scenario_rows
]

clone_table_rows = [
    [
        row["domain"],
        row["status"].replace("HTTP/2 ", ""),
        row["server"],
        row["x_powered_by"] or "-",
        "si" if row["wordpress_detected"] == "True" else "no",
        row["title"][:80],
    ]
    for row in clone_rows
]


md_lines = [
    "# Articulo tecnico: evidencias Futbol Libre y DoradoBet",
    "",
    "Fecha de cierre: 1 de junio de 2026. Evidencias recolectadas en entorno controlado entre el 31 de mayo y 1 de junio de 2026.",
    "",
    "## Tesis para el articulo",
    "",
    "No encontramos una descarga automatica de malware en las corridas realizadas. Lo peligroso aparece en otro lado: Futbol Libre combina streaming pirata, iframes, publicidad ofuscada y pop-under; DoradoBet concentra rastreo intensivo, fingerprinting, KYC y riesgo de clones/dominos parecidos.",
    "",
    "## Pruebas principales",
    "",
]

for proof in proofs:
    md_lines.extend(
        [
            f"### {proof['id']}. {proof['title']}",
            "",
            f"- Hallazgo: {proof['finding']}",
            f"- Evidencia: `{proof['evidence']}`",
            f"- Riesgo: {proof['risk']}",
            "",
        ]
    )

md_lines.extend(
    [
        "## Conteo por escenario",
        "",
        "| Escenario | Dominios | Popups | Descargas | localStorage |",
        "|---|---:|---:|---:|---:|",
    ]
)
for row in scenario_table_rows:
    md_lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} |")

md_lines.extend(
    [
        "",
        "## Dominios parecidos a DoradoBet",
        "",
        "| Dominio | Estado | Servidor | X-Powered-By | WordPress | Titulo |",
        "|---|---|---|---|---:|---|",
    ]
)
for row in clone_table_rows:
    md_lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} | {row[5]} |")

md_lines.extend(
    [
        "",
        "## Frases utilizables en el articulo",
        "",
        "- En nuestras pruebas no hubo descarga automatica, pero si hubo pop-under, iframes externos, publicidad ofuscada y rastreo de terceros.",
        "- En DoradoBet, el hallazgo mas fuerte no fue malware directo sino privacidad: 62 a 86 dominios externos por escenario y fingerprinting persistente.",
        "- En mobile/PWA aparecio una cadena de publicidad adulta/afiliada asociada a DoradoBet: bit.ly, go.xlviiirdr.com y chaturbate.com.",
        "- La superficie de riesgo no termina en el dominio oficial: existen dominios parecidos, clones o sitios promocionales que pueden capturar usuarios confundidos.",
        "",
        "## Archivos clave",
        "",
        "- `evidencias/publicable/ANEXO_EVIDENCIAS.md`",
        "- `evidencias/publicable/scenario_summary.csv`",
        "- `evidencias/publicable/domains_by_scenario.csv`",
        "- `evidencias/publicable/interesting_requests.csv`",
        "- `evidencias/publicable/doradobet_clone_scan.csv`",
        "- `evidencias/extended/doradobet_home_mobile/after.png`",
        "- `evidencias/extended/doradobet_register_desktop/after.png`",
        "- `evidencias/extended/futbol_espn_desktop_click/after.png`",
        "",
        "## Fuentes externas",
        "",
        "- https://doradobet.com/",
        "- https://verification.anjouangamingboard.org/",
        "- https://www.scamadviser.com/check-website-old/doradobet.com",
        "- https://www.scamdoc.com/view/419235",
        "- https://www.klazify.com/website/doradobet.com",
        "- https://www.yogonet.com/international/news/2024/04/16/71750-peru-doradobet-earns-authorization-to-operate-online-casino-games-for-the-next-six-years",
    ]
)

OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_MD.write_text("\n".join(md_lines) + "\n", encoding="utf-8")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
r = title.add_run("Articulo tecnico: evidencias Futbol Libre y DoradoBet")
r.font.name = "Calibri"
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
r = subtitle.add_run("Pruebas para sustentar un articulo de ciberseguridad")
r.font.name = "Calibri"
r.font.size = Pt(12)
r.italic = True

paragraph(doc, "Fecha de cierre: 1 de junio de 2026. Evidencias recolectadas en entorno controlado entre el 31 de mayo y 1 de junio de 2026.")
paragraph(doc, "Alcance: analisis defensivo/pasivo y navegacion controlada. No se crearon cuentas, no se realizaron pagos y no se ejecuto ningun archivo descargado.")

heading(doc, "Tesis para el articulo", 1)
paragraph(doc, "No encontramos una descarga automatica de malware en las corridas realizadas. Lo peligroso aparece en otro lado: Futbol Libre combina streaming pirata, iframes, publicidad ofuscada y pop-under; DoradoBet concentra rastreo intensivo, fingerprinting, KYC y riesgo de clones/dominios parecidos.")

heading(doc, "Pruebas principales", 1)
for proof in proofs:
    heading(doc, f"{proof['id']}. {proof['title']}", 2)
    paragraph(doc, f"Hallazgo: {proof['finding']}", bold_label="Hallazgo:")
    paragraph(doc, f"Evidencia: {proof['evidence']}", bold_label="Evidencia:")
    paragraph(doc, f"Riesgo: {proof['risk']}", bold_label="Riesgo:")

heading(doc, "Capturas clave", 1)
add_figure(
    doc,
    ROOT / "evidencias" / "extended" / "futbol_espn_desktop_click" / "after.png",
    "Figura 1. Futbol Libre, escenario ESPN despues de interaccion con el reproductor.",
)
add_figure(
    doc,
    ROOT / "evidencias" / "extended" / "doradobet_home_mobile" / "after.png",
    "Figura 2. DoradoBet mobile: portada usada para la prueba con mayor numero de dominios externos.",
)
add_figure(
    doc,
    ROOT / "evidencias" / "extended" / "doradobet_register_desktop" / "after.png",
    "Figura 3. DoradoBet registro: campos de datos personales/KYC visibles.",
)

heading(doc, "Conteo por escenario", 1)
table(
    doc,
    ["Escenario", "Dominios", "Popups", "Descargas", "localStorage"],
    scenario_table_rows,
    [2.7, 0.75, 0.75, 0.8, 1.0],
)

heading(doc, "Dominios parecidos a DoradoBet", 1)
table(
    doc,
    ["Dominio", "Estado", "Servidor", "X-Powered-By", "WP", "Titulo"],
    clone_table_rows,
    [1.35, 0.75, 0.8, 1.0, 0.45, 2.4],
)

heading(doc, "Frases utilizables en el articulo", 1)
for item in [
    "En nuestras pruebas no hubo descarga automatica, pero si hubo pop-under, iframes externos, publicidad ofuscada y rastreo de terceros.",
    "En DoradoBet, el hallazgo mas fuerte no fue malware directo sino privacidad: 62 a 86 dominios externos por escenario y fingerprinting persistente.",
    "En mobile/PWA aparecio una cadena de publicidad adulta/afiliada asociada a DoradoBet: bit.ly, go.xlviiirdr.com y chaturbate.com.",
    "La superficie de riesgo no termina en el dominio oficial: existen dominios parecidos, clones o sitios promocionales que pueden capturar usuarios confundidos.",
]:
    bullet(doc, item)

heading(doc, "Archivos clave", 1)
for item in [
    "evidencias/publicable/ANEXO_EVIDENCIAS.md",
    "evidencias/publicable/scenario_summary.csv",
    "evidencias/publicable/domains_by_scenario.csv",
    "evidencias/publicable/interesting_requests.csv",
    "evidencias/publicable/doradobet_clone_scan.csv",
    "evidencias/extended/doradobet_home_mobile/after.png",
    "evidencias/extended/doradobet_register_desktop/after.png",
    "evidencias/extended/futbol_espn_desktop_click/after.png",
]:
    bullet(doc, item)

heading(doc, "Fuentes externas", 1)
for item in [
    "https://doradobet.com/",
    "https://verification.anjouangamingboard.org/",
    "https://www.scamadviser.com/check-website-old/doradobet.com",
    "https://www.scamdoc.com/view/419235",
    "https://www.klazify.com/website/doradobet.com",
    "https://www.yogonet.com/international/news/2024/04/16/71750-peru-doradobet-earns-authorization-to-operate-online-casino-games-for-the-next-six-years",
]:
    bullet(doc, item)

doc.save(OUT_DOCX)

print(OUT_DOCX)
print(OUT_MD)
