from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Analisis_Ampliado_FutbolLibre_DoradoBet.docx"
OUT_MD = ROOT / "Analisis_Ampliado_FutbolLibre_DoradoBet.md"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Arial"
        p.add_run(text[len(bold_prefix):]).font.name = "Arial"
    else:
        p.add_run(text).font.name = "Arial"
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item).font.name = "Arial"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


report_md = """# Informe ampliado de ciberseguridad

Proyecto: Futbol Libre / DoradoBet  
Fecha local del análisis: 31 de mayo de 2026  
Autor: Walter Jaldin  
Objetivo: Profundizar en redirecciones, publicidad, descargas, rastreo y riesgos para el usuario.

## Resumen ejecutivo

El análisis dinámico de Futbol Libre confirmó que el mayor riesgo no aparece en la página inicial sino en la cadena reproductor-publicidad: el dominio original `futbollibretv.su` redirige a `futbol-libre.su`, los canales cargan iframes de `latamvidz1.com`, el reproductor integra publicidad tipo pop-under mediante `acscdn.com/script/aclib.js`, y el video usa infraestructura HLS/P2P con `envivoslatam.org` y `meshify.cloud`. En la prueba controlada no se observó descarga automática, pero sí se observó intento de abrir ventana nueva (`about:blank`) desde el iframe al interactuar con el reproductor.

DoradoBet presenta un perfil diferente. No se comportó como sitio pirata ni mostró descarga automática durante la prueba. Sí se identificó una superficie amplia de rastreo: FingerprintJS, Google Tag Manager, Amplitude, TikTok, Microsoft Clarity, Criteo, Adform, AppNexus/Xandr, MGID, Sportradar y otros dominios de publicidad/retargeting. El riesgo principal no es malware directo, sino privacidad, perfilado persistente, exposición financiera/KYC y suplantación por dominios parecidos.

## Metodología

- Reconocimiento DNS, TLS y cabeceras HTTP.
- Extracción de HTML, scripts, iframes y enlaces.
- Navegación automatizada con perfil temporal aislado de Chrome.
- Registro de solicitudes de red, dominios contactados, popups, service workers y descargas.
- No se creó cuenta, no se inició sesión, no se depositó dinero y no se ejecutó ningún archivo descargado.

## Futbol Libre: hallazgos nuevos

### Redirección principal

`https://futbollibretv.su/` redirige actualmente por HTTP 301 a `https://futbol-libre.su/`.

### Publicidad y scripts observados

En la portada se observaron:

- `https://acscdn.com/script/suv5.js`
- `https://acscdn.com/script/aclib.js`
- llamada explícita `aclib.runPop({ zoneId: '10652966' })`
- enlace publicitario oculto o sin texto visible hacia `https://adexchangerapid.com/ad/visit.php?al=1`
- Google Analytics / Google Tag Manager

### Canal ESPN / reproductor

La página `https://futbol-libre.su/espn-1/` incrusta:

- iframe: `https://latamvidz1.com/canal.php?stream=espn`
- reproductor Clappr
- plugin SwarmCloud HLS/P2P
- HLS servido desde `envivoslatam.org`
- comunicación P2P/telemetría con `us.meshify.cloud`
- service worker en `https://latamvidz1.com/sw.js`

### Comportamiento al interactuar

Durante la prueba controlada:

- No se observó descarga automática.
- No se observó archivo `.exe`, `.apk`, `.dmg`, `.zip` ni equivalente.
- Se observó un intento de abrir ventana nueva desde el iframe (`Page.windowOpen` con `about:blank`) al hacer clic en el reproductor.
- La publicidad se ejecuta dentro de un contexto de terceros y puede cambiar por país, hora, IP, navegador y campaña.

### Riesgo técnico

El riesgo se mantiene alto porque la página combina contenido pirata, publicidad de baja calidad, scripts ofuscados, iframes externos, pop-under y P2P. Aunque en esta ejecución no se descargó malware, el usuario queda expuesto a redirecciones variables, rastreo, abuso de permisos, ventanas emergentes y campañas de malvertising que pueden cambiar sin modificar el dominio principal.

## DoradoBet: hallazgos técnicos

### Dominio analizado

Dominio principal observado: `doradobet.com`

DNS:

- IPv4: `104.18.192.230`, `104.18.193.230`
- IPv6: `2606:4700::6812:c0e6`, `2606:4700::6812:c1e6`
- Nameservers: `edna.ns.cloudflare.com`, `jerry.ns.cloudflare.com`
- Proveedor frontal: Cloudflare

TLS:

- Certificado para `doradobet.com` y `*.doradobet.com`
- Emisor: Google Trust Services
- Vigencia observada: 16 de mayo de 2026 a 14 de agosto de 2026

Cabeceras positivas:

- Redirección HTTP a HTTPS
- `Strict-Transport-Security`
- `X-Content-Type-Options: nosniff`
- cookies Cloudflare y balanceador con `HttpOnly` / `Secure`

Cabeceras débiles o incompletas:

- CSP observada solo restringe `frame-ancestors`; no se observó una política fuerte que limite scripts, conexiones e imágenes.
- No se observó `Permissions-Policy`.
- `Referrer-Policy` es `no-referrer-when-downgrade`, menos restrictiva que `strict-origin-when-cross-origin` o `no-referrer`.
- No se encontró `security.txt`; la URL respondió con HTML normal.

### Rastreo y terceros

En una carga controlada de `https://doradobet.com/` se observaron más de 50 dominios externos. Entre los principales:

- `fpjscdn.net` y `api.fpjs.io` para FingerprintJS
- `static.virtualsoft.tech` y `cdnconfigs.virtualsoft.tech`
- `api2.amplitude.com` y `cdn.amplitude.com`
- `www.googletagmanager.com`, `www.google-analytics.com`, `stats.g.doubleclick.net`
- `analytics.tiktok.com`
- `www.clarity.ms`, `scripts.clarity.ms`
- `dynamic.criteo.com`, `gum.criteo.com`
- `a2.adform.net`, `track.adform.net`
- `acdn.adnxs.com`, `ib.adnxs.com`
- `a.mgid.com`
- `tracker.ads.sportradar.com`, `tm.ads.sportradar.com`
- `ctag.containermedia.net`
- `www.rtb123.com`, `x.bidswitch.net`, `pixel.rubiconproject.com`

### Fingerprinting confirmado

El HTML contiene una integración explícita con FingerprintJS:

- carga dinámica de `https://fpjscdn.net/v4/PmGJOlgFXlEl6IqCC6Bi`
- almacenamiento de `result.visitor_id` en `localStorage`
- claves observadas: `visitor_fp`, `visitor_fp5`, `fp_visitor_id`
- envío del identificador a `https://static.virtualsoft.tech/setvid`

Esto permite reconocer al visitante de manera persistente incluso si no inicia sesión. No equivale a malware, pero sí es rastreo avanzado.

### Permisos, descargas y popups

Durante la prueba:

- No se observó descarga automática.
- No se observaron instaladores `.apk`, `.exe`, `.dmg`, `.pkg` o `.zip`.
- No se observó popup en la carga inicial.
- No se observó solicitud de notificaciones push; el permiso quedó en `default`.
- No se observó service worker activo en la página principal.
- La sección de app apunta a una PWA (`/landing/app-pwa`), no a una descarga directa de APK en la evidencia recolectada.

### Licencia y jurisdicción

La propia página muestra en el pie que DoradoBet es operado por `VS Services Ltd`, con dirección en Anjouan, Unión de Comoras, y licencia `ALSI-132405045-F13` o variante visualmente similar. El sello externo de Anjouan confirmó una página de validación para `doradobet.com` y mostró licencia `ALSI-132405045-FI3`, con jurisdicciones excluidas como Estados Unidos, Reino Unido, España, Francia, Países Bajos, Australia y Comoros.

Este punto no implica malware, pero sí es importante para el usuario: cualquier disputa de cuenta, retiro, bloqueo o verificación KYC dependerá de términos y jurisdicción de juego, no de protecciones bancarias comunes.

### Riesgos para el usuario

- Privacidad: alto. El sitio usa fingerprinting persistente y muchas redes de publicidad/retargeting.
- Malware directo: bajo en esta ejecución. No hubo descarga ni exploit observado.
- Phishing/suplantación: medio-alto. Hay múltiples dominios parecidos en resultados públicos (`doradobet.online`, `dorado-bet.pe`, `doradobet-peru.org`, etc.). Algunos son WordPress o sitios promocionales no necesariamente oficiales.
- Riesgo financiero: alto por naturaleza del servicio. El usuario puede entregar datos personales, documento, teléfono, banco y medios de pago.
- Riesgo KYC: alto. Para retiros/verificación se solicitan datos reales y documentos; si el usuario cae en un clon, el daño potencial es robo de identidad.
- Riesgo legal/regulatorio: variable por país. El usuario debe verificar si el operador está autorizado en su jurisdicción concreta.

## Comparación rápida

| Sitio | Riesgo principal | Descarga observada | Publicidad/redirección | Rastreo |
|---|---|---:|---|---|
| Futbol Libre | Malvertising, pop-under, iframes externos, P2P | No | Alta | Media |
| DoradoBet | Privacidad, KYC, financiero, suplantación por clones | No | Alta red publicitaria/retargeting, no popup inicial | Alta |

## Conclusiones

No se encontró una descarga maliciosa automática ni en Futbol Libre ni en DoradoBet durante estas ejecuciones controladas. Sin embargo, eso no significa que sean equivalentes.

Futbol Libre sigue siendo el más peligroso técnicamente: usa infraestructura cambiante de streaming pirata, publicidad pop-under y terceros que pueden variar por campaña. La ausencia de descarga en una prueba no elimina el riesgo de malvertising.

DoradoBet parece una plataforma comercial más estructurada y con controles básicos de seguridad web, pero expone al usuario a rastreo intensivo, perfilado persistente, recopilación KYC y riesgo financiero. El mayor peligro práctico es entregar datos reales o documentos a un clon o a una página promocional no oficial.

## Recomendaciones

- No usar perfiles personales del navegador para sitios de apuestas o streaming riesgoso.
- Bloquear terceros con uBlock Origin o Brave Shields.
- Revisar y borrar permisos del navegador, cookies y datos de sitio.
- No instalar APKs ni apps fuera de tiendas oficiales.
- Verificar que el dominio sea exactamente `doradobet.com` antes de ingresar datos.
- No enviar documentos por correo o chat salvo que se haya validado el canal oficial.
- Usar tarjetas virtuales o métodos con límites cuando se trate de pagos en línea.
- Revisar términos de bonos, rollover, KYC, retiros y jurisdicción antes de depositar.

## Evidencias locales

- `evidencias/aclib.js`
- `evidencias/suv5.js`
- `evidencias/latamvidz1_espn.html`
- `evidencias/cdp/probe-result.json`
- `evidencias/doradobet/home.html`
- `evidencias/doradobet/pe_home.html`
- `evidencias/doradobet/app_pwa.html`
- `evidencias/doradobet/config_bo.js`
- `evidencias/doradobet/cdp/probe-result.json`

## Fuentes consultadas

- Sitio principal: https://doradobet.com/
- Validación Anjouan: https://verification.anjouangamingboard.org/
- ScamAdviser: https://www.scamadviser.com/check-website-old/doradobet.com
- ScamDoc: https://www.scamdoc.com/view/419235
- Klazify: https://www.klazify.com/website/doradobet.com
- Yogonet sobre autorización en Perú: https://www.yogonet.com/international/news/2024/04/16/71750-peru-doradobet-earns-authorization-to-operate-online-casino-games-for-the-next-six-years
"""


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informe ampliado de ciberseguridad")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Futbol Libre / DoradoBet - 31 de mayo de 2026")
    r.font.name = "Arial"
    r.font.size = Pt(11)

    add_heading(doc, "Resumen ejecutivo", 1)
    add_para(doc, "El análisis dinámico de Futbol Libre confirmó que el mayor riesgo no aparece en la página inicial sino en la cadena reproductor-publicidad: redirección a futbol-libre.su, iframes de latamvidz1.com, pop-under mediante acscdn.com y video HLS/P2P con envivoslatam.org y meshify.cloud. No se observó descarga automática, pero sí intento de abrir ventana nueva desde el iframe.")
    add_para(doc, "DoradoBet presentó un perfil diferente: no hubo descarga automática ni popup inicial, pero sí rastreo intensivo, FingerprintJS, múltiples redes de publicidad/retargeting y almacenamiento de identificadores persistentes. Su riesgo principal es privacidad, KYC, financiero y suplantación por dominios parecidos.")

    add_heading(doc, "Metodología", 1)
    add_bullets(doc, [
        "Reconocimiento DNS, TLS y cabeceras HTTP.",
        "Extracción de HTML, scripts, iframes y enlaces.",
        "Navegación automatizada con perfil temporal aislado de Chrome.",
        "Registro de solicitudes de red, dominios contactados, popups, service workers y descargas.",
        "No se creó cuenta, no se inició sesión, no se depositó dinero y no se ejecutó ningún archivo descargado.",
    ])

    add_heading(doc, "Futbol Libre: hallazgos nuevos", 1)
    add_bullets(doc, [
        "futbollibretv.su redirige actualmente a futbol-libre.su.",
        "La portada carga acscdn.com/script/suv5.js y acscdn.com/script/aclib.js.",
        "Se observó aclib.runPop con zoneId 10652966 y un enlace publicitario hacia adexchangerapid.com.",
        "El canal ESPN carga un iframe de latamvidz1.com con reproductor Clappr, SwarmCloud HLS/P2P y service worker.",
        "La reproducción usa segmentos HLS desde envivoslatam.org y comunicación con us.meshify.cloud.",
        "No se observó descarga automática; sí se observó intento de abrir ventana nueva desde el iframe.",
    ])

    add_heading(doc, "DoradoBet: hallazgos técnicos", 1)
    add_table(doc, ["Área", "Hallazgo"], [
        ["Dominio", "doradobet.com"],
        ["Infraestructura", "Cloudflare; A 104.18.192.230 / 104.18.193.230"],
        ["TLS", "Google Trust Services; wildcard *.doradobet.com"],
        ["Cabeceras positivas", "HTTPS, HSTS, nosniff, cookies Secure/HttpOnly"],
        ["Debilidades", "CSP limitada a frame-ancestors; sin Permissions-Policy; Referrer-Policy permisiva"],
        ["Descargas", "No se observaron descargas automáticas ni instaladores"],
        ["Notificaciones", "No se observó solicitud push; permiso quedó en default"],
        ["PWA/App", "Sección app-pwa; no APK observado en la evidencia"],
    ])

    add_heading(doc, "Rastreo y terceros", 2)
    add_para(doc, "En una sola carga se observaron más de 50 dominios externos de analítica, publicidad y retargeting. Entre ellos: FingerprintJS, Amplitude, Google Tag Manager/Analytics, TikTok, Microsoft Clarity, Criteo, Adform, AppNexus/Xandr, MGID, Sportradar, RTB123, Bidswitch y Rubicon.")

    add_heading(doc, "Fingerprinting confirmado", 2)
    add_bullets(doc, [
        "Carga dinámica de https://fpjscdn.net/v4/PmGJOlgFXlEl6IqCC6Bi.",
        "Almacena visitor_id en localStorage como visitor_fp, visitor_fp5 y fp_visitor_id.",
        "Envía el identificador a https://static.virtualsoft.tech/setvid.",
        "Esto permite reconocer al visitante de forma persistente, incluso sin iniciar sesión.",
    ])

    add_heading(doc, "Licencia y jurisdicción", 2)
    add_para(doc, "La página muestra operación por VS Services Ltd, dirección en Anjouan, Unión de Comoros, y licencia ALSI-132405045-F13 o variante visualmente similar. El sello externo de Anjouan validó el dominio doradobet.com y muestra licencia ALSI-132405045-FI3. La validación indica jurisdicciones excluidas como Estados Unidos, Reino Unido, España, Francia, Países Bajos, Australia y Comoros.")

    add_heading(doc, "Comparación rápida", 1)
    add_table(doc, ["Sitio", "Riesgo principal", "Descarga", "Rastreo"], [
        ["Futbol Libre", "Malvertising, pop-under, iframes externos, P2P", "No observada", "Medio"],
        ["DoradoBet", "Privacidad, KYC, financiero, suplantación por clones", "No observada", "Alto"],
    ])

    add_heading(doc, "Conclusiones", 1)
    add_para(doc, "No se encontró descarga maliciosa automática ni en Futbol Libre ni en DoradoBet durante estas ejecuciones controladas. Futbol Libre sigue siendo más peligroso técnicamente por publicidad ofuscada, pop-under, streaming pirata e infraestructura cambiante.")
    add_para(doc, "DoradoBet parece una plataforma comercial más estructurada y con controles básicos de seguridad web, pero expone al usuario a rastreo intensivo, perfilado persistente, recopilación KYC y riesgo financiero. El mayor peligro práctico es entregar datos reales o documentos a un clon o a una página promocional no oficial.")

    add_heading(doc, "Recomendaciones", 1)
    add_bullets(doc, [
        "No usar el perfil personal del navegador en sitios de riesgo.",
        "Usar bloqueador de terceros y revisar permisos/cookies después de la visita.",
        "Verificar que el dominio sea exactamente doradobet.com antes de ingresar datos.",
        "No instalar APKs ni apps fuera de tiendas oficiales.",
        "No enviar documentos por correo/chat sin validar el canal oficial.",
        "Leer condiciones de bonos, rollover, KYC, retiros y jurisdicción antes de depositar.",
    ])

    add_heading(doc, "Evidencias locales", 1)
    add_bullets(doc, [
        "evidencias/cdp/probe-result.json",
        "evidencias/latamvidz1_espn.html",
        "evidencias/aclib.js y evidencias/suv5.js",
        "evidencias/doradobet/cdp/probe-result.json",
        "evidencias/doradobet/home.html",
        "evidencias/doradobet/app_pwa.html",
        "evidencias/doradobet/config_bo.js",
    ])

    add_heading(doc, "Fuentes consultadas", 1)
    add_bullets(doc, [
        "https://doradobet.com/",
        "https://verification.anjouangamingboard.org/",
        "https://www.scamadviser.com/check-website-old/doradobet.com",
        "https://www.scamdoc.com/view/419235",
        "https://www.klazify.com/website/doradobet.com",
        "https://www.yogonet.com/international/news/2024/04/16/71750-peru-doradobet-earns-authorization-to-operate-online-casino-games-for-the-next-six-years",
    ])

    doc.save(OUT_DOCX)


def main():
    OUT_MD.write_text(report_md, encoding="utf-8")
    build_docx()
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
